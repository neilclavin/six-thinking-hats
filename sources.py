"""Source ingestion for the Six Thinking Hats council.

Loaders pull material from local folders and the web and yield uniform
``Document`` objects. The council then reasons over those documents (see
``corpus.build_corpus``) instead of a bare problem string.

Design: "files & folders only" — no API keys or OAuth. Everything here runs
against local files and plain HTTP. Live connectors (Gmail/Drive/Slack) and
Figma are intentionally left as documented stubs.
"""

from __future__ import annotations

import base64
import glob as globmod
import os
import urllib.request
from dataclasses import dataclass
from html.parser import HTMLParser
from typing import Iterator

# Extension → how to read it.
_TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".log", ".json", ".yaml", ".yml"}
_HTML_EXTS = {".html", ".htm"}
_IMAGE_EXTS = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}


@dataclass
class Document:
    """One unit of source material.

    Exactly one of ``text`` or ``image`` is populated. Image documents are read
    by the vision model during the digest step and turned into text.
    """

    title: str
    source_type: str  # "file" | "web"
    uri: str
    text: str | None = None
    image: tuple[str, str] | None = None  # (media_type, base64_data)

    @property
    def is_image(self) -> bool:
        return self.image is not None


# --------------------------------------------------------------------------- #
# HTML → text (stdlib only, no bs4)
# --------------------------------------------------------------------------- #
class _TextExtractor(HTMLParser):
    _SKIP = {"script", "style", "head", "noscript", "svg"}

    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skipping = 0

    def handle_starttag(self, tag: str, attrs: object) -> None:
        if tag in self._SKIP:
            self._skipping += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._skipping:
            self._skipping -= 1

    def handle_data(self, data: str) -> None:
        if not self._skipping and data.strip():
            self._chunks.append(data.strip())

    def text(self) -> str:
        return "\n".join(self._chunks)


def _html_to_text(html: str) -> str:
    parser = _TextExtractor()
    parser.feed(html)
    return parser.text()


# --------------------------------------------------------------------------- #
# Local folder
# --------------------------------------------------------------------------- #
class LocalFolderSource:
    """Walk a folder and yield a ``Document`` per supported file.

    Supported: text (.txt/.md/.csv/.log/.json/.yaml), HTML, PDF (needs
    ``pypdf``), DOCX (needs ``python-docx``), images (.png/.jpg/.gif/.webp).
    Unsupported extensions are skipped with a note on stderr.
    """

    def __init__(self, path: str, recursive: bool = True) -> None:
        self.path = path
        self.recursive = recursive

    def load(self) -> Iterator[Document]:
        if not os.path.isdir(self.path):
            raise NotADirectoryError(f"--sources path is not a folder: {self.path}")
        pattern = "**/*" if self.recursive else "*"
        for fp in sorted(globmod.glob(os.path.join(self.path, pattern), recursive=self.recursive)):
            if not os.path.isfile(fp):
                continue
            doc = self._load_file(fp)
            if doc is not None:
                yield doc

    def _load_file(self, fp: str) -> Document | None:
        ext = os.path.splitext(fp)[1].lower()
        title = os.path.relpath(fp, self.path)
        try:
            if ext in _TEXT_EXTS:
                return Document(title, "file", fp, text=_read_text(fp))
            if ext in _HTML_EXTS:
                return Document(title, "file", fp, text=_html_to_text(_read_text(fp)))
            if ext in _PDF_EXTS:
                return Document(title, "file", fp, text=_read_pdf(fp))
            if ext in _DOCX_EXTS:
                return Document(title, "file", fp, text=_read_docx(fp))
            if ext in _IMAGE_EXTS:
                with open(fp, "rb") as f:
                    data = base64.standard_b64encode(f.read()).decode("ascii")
                return Document(title, "file", fp, image=(_IMAGE_EXTS[ext], data))
        except Exception as exc:  # keep going; one bad file shouldn't kill the run
            _warn(f"skipping {title}: {exc}")
            return None
        _warn(f"skipping {title}: unsupported file type {ext!r}")
        return None


def _read_text(fp: str) -> str:
    with open(fp, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _read_pdf(fp: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # noqa: F841
        raise RuntimeError("PDF support needs pypdf — run: pip install pypdf") from None
    reader = PdfReader(fp)
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _read_docx(fp: str) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise RuntimeError("DOCX support needs python-docx — run: pip install python-docx") from None
    document = docx.Document(fp)
    return "\n".join(p.text for p in document.paragraphs).strip()


# --------------------------------------------------------------------------- #
# Web
# --------------------------------------------------------------------------- #
class WebSource:
    """Fetch a list of URLs over HTTP and yield text ``Document``s."""

    _UA = "Mozilla/5.0 (compatible; SixHatsBot/1.0; +https://example.com/bot)"

    def __init__(self, urls: list[str], timeout: float = 30.0) -> None:
        self.urls = urls
        self.timeout = timeout

    def load(self) -> Iterator[Document]:
        for url in self.urls:
            try:
                req = urllib.request.Request(url, headers={"User-Agent": self._UA})
                with urllib.request.urlopen(req, timeout=self.timeout) as resp:
                    ctype = resp.headers.get_content_type()
                    raw = resp.read().decode("utf-8", errors="replace")
                text = _html_to_text(raw) if "html" in ctype else raw
            except Exception as exc:
                _warn(f"skipping url {url}: {exc}")
                continue
            yield Document(url, "web", url, text=text)


def read_url_list(path: str) -> list[str]:
    """Read one URL per line from a file (blank lines and #comments ignored)."""
    urls: list[str] = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#"):
                urls.append(line)
    return urls


# --------------------------------------------------------------------------- #
# Figma — deferred stub
# --------------------------------------------------------------------------- #
class FigmaSource:
    """Placeholder for Figma ingestion (not implemented).

    Figma has no plain-text export and its pages don't fetch as HTML text, so
    it doesn't fit the "files & folders only" model without extra work. Two
    future paths:

    1. Frame export — export the relevant frames/pages as PNG/PDF/SVG into the
       source folder; ``LocalFolderSource`` then handles them (images are read
       by the vision model during digest). Zero keys.
    2. Figma REST API — set a read-only ``FIGMA_TOKEN`` and pass a file key/URL;
       pull text layers + node metadata via
       ``GET https://api.figma.com/v1/files/{key}`` and optionally rendered
       images via the ``/images`` endpoint.
    """

    def load(self) -> Iterator[Document]:
        raise NotImplementedError(
            "FigmaSource is a stub. For now, export Figma frames as PNG/PDF into "
            "your --sources folder, or implement the FIGMA_TOKEN REST path."
        )


# --------------------------------------------------------------------------- #
def _warn(msg: str) -> None:
    import sys

    print(f"[sources] {msg}", file=sys.stderr)


def collect(
    folders: list[str] | None = None,
    urls: list[str] | None = None,
) -> list[Document]:
    """Convenience: run the file and web loaders and return all documents."""
    docs: list[Document] = []
    for folder in folders or []:
        docs.extend(LocalFolderSource(folder).load())
    if urls:
        docs.extend(WebSource(urls).load())
    return docs
